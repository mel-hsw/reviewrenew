"""
Build the final project report as a Word .docx file.
Run from the project root:
    python scripts/build_report_docx.py
Outputs: docs/report_final.docx
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE = Path(__file__).resolve().parent.parent
OUTPUTS = BASE / "outputs"
DOCS = BASE / "docs"
DOCS.mkdir(exist_ok=True)

IMG = {
    # Pet bed — B0C4BFWLNC
    "bed_hero_oai":     OUTPUTS / "B0C4BFWLNC/openai/00_1_hero_product_chatgpt_image_latest_00.png",
    "bed_hero_gem":     OUTPUTS / "B0C4BFWLNC/gemini/00_1_hero_product_gemini-3_1-flash-image-preview_00.png",
    "bed_use_oai":      OUTPUTS / "B0C4BFWLNC/openai/01_2_consumer_in_use_chatgpt_image_latest_00.png",
    "bed_use_gem":      OUTPUTS / "B0C4BFWLNC/gemini/01_2_consumer_in_use_gemini-3_1-flash-image-preview_00.png",
    "bed_bundle_oai":   OUTPUTS / "B0C4BFWLNC/openai/02_3_wash_and_bundle_value_chatgpt_image_latest_00.png",
    "bed_bundle_gem":   OUTPUTS / "B0C4BFWLNC/gemini/02_3_wash_and_bundle_value_gemini-3_1-flash-image-preview_00.png",
    "bed_detail_oai":   OUTPUTS / "B0C4BFWLNC/openai/02_3_construction_detail_chatgpt_image_latest_00.png",
    "bed_detail_gem":   OUTPUTS / "B0C4BFWLNC/gemini/02_3_construction_detail_gemini-3_1-flash-image-preview_00.png",
    # CeraVe — B0BG9Q18ZZ
    "cer_hero_oai":     OUTPUTS / "B0BG9Q18ZZ/openai/00_1_hero_product_chatgpt_image_latest_00.png",
    "cer_hero_gem":     OUTPUTS / "B0BG9Q18ZZ/gemini/00_1_hero_product_gemini-3_1-flash-image-preview_00.png",
    "cer_use_oai":      OUTPUTS / "B0BG9Q18ZZ/openai/01_2_consumer_in_use_chatgpt_image_latest_00.png",
    "cer_use_gem":      OUTPUTS / "B0BG9Q18ZZ/gemini/01_2_consumer_in_use_gemini-3_1-flash-image-preview_00.png",
    "cer_tex_oai":      OUTPUTS / "B0BG9Q18ZZ/openai/02_3_texture_formula_closeup_chatgpt_image_latest_00.png",
    "cer_tex_gem":      OUTPUTS / "B0BG9Q18ZZ/gemini/02_3_texture_formula_closeup_gemini-3_1-flash-image-preview_00.png",
    # Hat washer — B0BXSQN4HV
    "hat_hero_oai":     OUTPUTS / "B0BXSQN4HV/openai/00_1_hero_product_chatgpt_image_latest_00.png",
    "hat_hero_gem":     OUTPUTS / "B0BXSQN4HV/gemini/00_1_hero_product_gemini-3_1-flash-image-preview_00.png",
    "hat_use_oai":      OUTPUTS / "B0BXSQN4HV/openai/01_2_consumer_in_use_chatgpt_image_latest_00.png",
    "hat_use_gem":      OUTPUTS / "B0BXSQN4HV/gemini/01_2_consumer_in_use_gemini-3_1-flash-image-preview_00.png",
    "hat_mac_oai":      OUTPUTS / "B0BXSQN4HV/openai/02_3_detail_macro_chatgpt_image_latest_00.png",
    "hat_mac_gem":      OUTPUTS / "B0BXSQN4HV/gemini/02_3_detail_macro_gemini-3_1-flash-image-preview_00.png",
}

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    run = p.runs[0]
    run.italic = True
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_cell_border(cell, **kwargs) -> None:
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"), kwargs.get(side, "none"))
        tag.set(qn("w:sz"), "0")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "auto")
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def side_by_side(
    doc: Document,
    left_path: Path,
    right_path: Path,
    left_label: str,
    right_label: str,
    fig_caption: str,
    img_width: float = 2.9,
) -> None:
    """Insert two images side-by-side with labels, in a borderless table."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"

    # Label row
    for i, lbl in enumerate((left_label, right_label)):
        cell = tbl.cell(0, i)
        cell.text = lbl
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_border(cell, top="none", left="none", bottom="none", right="none")

    # Image row
    for i, img_path in enumerate((left_path, right_path)):
        cell = tbl.cell(1, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        if img_path.exists():
            run.add_picture(str(img_path), width=Inches(img_width))
        else:
            cell.text = f"[image not found: {img_path.name}]"
        set_cell_border(cell, top="none", left="none", bottom="none", right="none")

    caption(doc, fig_caption)
    doc.add_paragraph()


def single_image(doc: Document, img_path: Path, fig_caption: str, width: float = 4.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if img_path.exists():
        run.add_picture(str(img_path), width=Inches(width))
    else:
        p.text = f"[image not found: {img_path.name}]"
    caption(doc, fig_caption)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # Header
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        hdr.cells[i].paragraphs[0].runs[0].bold = True
        hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    # Data rows
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            tbl.rows[ri].cells[ci].text = val
            tbl.rows[ri].cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build() -> None:
    doc = Document()

    # Page margins (1 inch all round)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.15)
        section.right_margin = Inches(1.15)

    # Normal style font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Generating Product Images from Customer Reviews")
    title_run.bold = True
    title_run.font.size = Pt(20)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Final Project Report\n").font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        "94-844 Generative AI Lab · Spring 2026\n"
        "Prof. Beibei Li · Heinz College · Carnegie Mellon University\n"
        "Due: April 26, 2026"
    )
    meta_run.font.size = Pt(12)

    add_page_break(doc)

    # =========================================================================
    # ABSTRACT
    # =========================================================================
    heading(doc, "Abstract", 1)
    body(doc,
        "This project builds an end-to-end pipeline for generating product marketing imagery "
        "grounded in customer review text. We selected three Amazon products spanning Pet Supplies "
        "(Best Friends by Sheri calming donut bed bundle, B0C4BFWLNC), Beauty & Personal Care "
        "(CeraVe Hydrating Facial Cleanser, B0BG9Q18ZZ), and Home laundry accessories (YEENOR hat "
        "washer cage 2-pack, B0BXSQN4HV). These categories were chosen to represent soft-goods "
        "texture, labeled CPG packaging, and rigid-plastic utility — each presenting distinct "
        "challenges for text-to-image generation."
    )
    body(doc,
        "The pipeline chunks listing text and reviews using tiktoken, embeds them in a Chroma "
        "vector store, retrieves relevant excerpts with 12-query multi-query RAG, and prompts a "
        "structured JSON LLM (gpt-5.4-mini) to produce an ImagePromptBundle: a typed schema "
        "containing review-grounded summaries, brand visual codes, and 3–5 planned shot prompts "
        "per product. Images are generated with two APIs — OpenAI (chatgpt-image-latest) and "
        "Google Gemini (gemini-3.1-flash-image-preview) — under the same base prompts plus "
        "provider-specific suffixes, yielding 18 images total across the three products."
    )
    body(doc,
        "We document three deliberate experiments: a chunking A/B comparison (sliding windows vs. "
        "paragraph-batch packing), a dual-model image generation study (same prompts, two backends), "
        "and an iterative prompt engineering study informed by inspecting real listing_image_urls "
        "and observed diffusion failure modes. Key findings: paragraph-batch chunking better matches "
        "short Amazon reviews; OpenAI images lean toward tighter catalog/studio frames while Gemini "
        "adds more lifestyle context; and neither model reliably reproduces certification micro-copy "
        "or dense label panels, requiring prompt constraints that favour silhouette, palette, and "
        "short verified phrases."
    )

    add_page_break(doc)

    # =========================================================================
    # Q1 — PRODUCT SELECTION
    # =========================================================================
    heading(doc, "Q1 — Product Selection and Data Collection", 1)

    heading(doc, "1.1  Selection Rationale", 2)
    body(doc,
        "The assignment requires three products from different categories. Beyond category "
        "diversity, we applied a backward-design criterion: each product needed to create "
        "meaningfully different challenges for the Q2 → Q3 pipeline so that the Q3 comparisons "
        "would be analytically interesting rather than redundant."
    )

    add_table(doc,
        headers=["Criterion", "Pet Bed (B0C4BFWLNC)", "CeraVe Cleanser (B0BG9Q18ZZ)", "Hat Washer (B0BXSQN4HV)"],
        rows=[
            ["Category", "Pet Supplies", "Beauty & Personal Care", "Home / Laundry Accessories"],
            ["Primary visual challenge", "Soft-goods texture; 'does it really look plush?'",
             "Labeled CPG bottle; text-on-pack, non-foaming consistency",
             "Rigid plastic; unusual 3-D cage geometry"],
            ["Key review language", "Shape, loft, washability, colour accuracy",
             "Non-foaming lotion texture, barrier claims, packaging leaks",
             "Clip security, shape retention, brim fit"],
            ["Why this fuels Q3", "Risk: flat/misleading staging of fur and loft",
             "Diffusion struggles with legible micro-text on label",
             "Geometry test: can diffusion render a molded plastic cage?"],
            ["Review volume (meta)", "~68,128 ratings", "Thousands (up to 500 used)", "~21,076 ratings"],
        ]
    )

    body(doc,
        "Trio rationale: The three products span soft texture (bed), consumable with pharmaceutical-"
        "style labeling (cleanser), and hard-surface utility (hat washer). They produce different "
        "failure modes in Q3 — plushness hallucinations vs. text hallucinations vs. geometry "
        "distortion — making the cross-product comparison analytically informative."
    )

    heading(doc, "1.2  Products", 2)

    heading(doc, "Product 1 — Best Friends by Sheri Calming Lux Donut Cuddler + Blanket Bundle", 3)
    add_table(doc,
        headers=["Field", "Value"],
        rows=[
            ["ASIN", "B0C4BFWLNC"],
            ["Category", "Pet Supplies — Calming Beds"],
            ["Title", "Best Friends by Sheri Bundle Set The Original Calming Lux Donut Cuddler Cat and Dog Bed + Pet Throw Blanket Dark Chocolate Medium 30\" × 30\""],
            ["URL", "https://www.amazon.com/dp/B0C4BFWLNC"],
            ["Reviews used", "Up to 500, Amazon Reviews 2023 Pet Supplies JSONL, newest-first, len > 20"],
        ]
    )
    body(doc,
        "Reviews extensively discuss whether the bed looks as plush in person as in photos, colour "
        "accuracy ('Dark Chocolate'), washability after cycles, and whether the rim stays lofted. "
        "These are precisely the visual trust signals an AI image system must get right — or risk "
        "compounding the 'misleading photo' problem reviewers already flag."
    )

    heading(doc, "Product 2 — CeraVe Hydrating Facial Cleanser (16 fl oz)", 3)
    add_table(doc,
        headers=["Field", "Value"],
        rows=[
            ["ASIN", "B0BG9Q18ZZ"],
            ["Category", "Beauty & Personal Care — Skin Care — Face — Cleansers"],
            ["Title", "CeraVe Hydrating Facial Cleanser | Moisturizing Non-Foaming Face Wash with Hyaluronic Acid, Ceramides and Glycerin | Fragrance Free | 16 Fluid Ounce"],
            ["URL", "https://www.amazon.com/dp/B0BG9Q18ZZ"],
            ["Reviews used", "Up to 500, Beauty JSONL slice, same filter"],
        ]
    )
    body(doc,
        "CeraVe is a high-traffic, trust-heavy skincare SKU with rich review language around "
        "texture (non-foaming, lotion-like), skin feel (hydrated, not stripped), fragrance-free "
        "sensitivity, and packaging reliability (pump leaks). The label-heavy bottle creates a "
        "classic diffusion challenge: models attempt to render ingredient panels and seals, "
        "producing plausible-but-wrong glyphs."
    )

    heading(doc, "Product 3 — YEENOR Hat Washer for Washing Machine, 2-Pack", 3)
    add_table(doc,
        headers=["Field", "Value"],
        rows=[
            ["ASIN", "B0BXSQN4HV"],
            ["Category", "Home / Laundry Accessories"],
            ["Title", "YEENOR Hat Washer for Washing Machine, 2-Pack Baseball Caps Washers Hat Storage, Hat Holder Fit for Adult/Kid's Hat"],
            ["URL", "https://www.amazon.com/dp/B0BXSQN4HV"],
            ["Reviews used", "Up to 500, Home JSONL slice, same filter"],
        ]
    )
    body(doc,
        "A molded plastic cage with a specific 3-D silhouette (four-sided, ribbed, clip-secured) "
        "is an interesting geometry test for diffusion. Reviews concentrate on shape retention, "
        "clip security through the wash cycle, brim fit, and plastic durability — all functional "
        "and structural, making it a strong contrast to the soft-goods bed and the cosmetic cleanser."
    )

    heading(doc, "1.3  Data Collection Protocol", 2)

    body(doc,
        "Attempted direct scraping — blocked. Our first approach was to scrape listing text and "
        "reviews directly from Amazon product pages using the fixed ASINs. Amazon's bot-detection "
        "systems returned 503 errors and CAPTCHA challenges consistently, making live scraping "
        "unreliable for a reproducible academic pipeline. We therefore pivoted to an established, "
        "citable academic dataset.\n\n"
        "Final source — Amazon Reviews 2023 (Hou et al., 2024). Review and metadata slices come "
        "from the Amazon Reviews 2023 dataset (amazon-reviews-2023.github.io). The dataset provides "
        "per-category JSONL files containing product metadata (title, description, feature bullets, "
        "listing image URLs, category breadcrumbs) and user reviews (text, rating, date). We "
        "downloaded the relevant category slices — Pet Supplies, Beauty & Personal Care, and "
        "Appliances/Home — and filtered them to the three target ASINs using "
        "scripts/consolidate_products.py and scripts/filter_meta_by_rating.py."
    )

    add_table(doc,
        headers=["Item", "Detail"],
        rows=[
            ["Primary source", "Amazon Reviews 2023 (Hou et al., 2024; McAuley Lab, UCSD) — JSONL per category"],
            ["Data access method", "Pre-collected academic dataset; direct Amazon scraping was blocked by bot-detection"],
            ["Review filter", "len(text) > 20; newest first; cap at 500 reviews per product"],
            ["Fields captured", "text, rating, title, date (review metadata)"],
            ["Listing fields", "title, description (features/bullets), listing_image_urls (carousel order), category"],
            ["Consolidation script", "scripts/consolidate_products.py → data/products.json + data/products_manifest.json"],
            ["Collection date", "2026-04-18 (pipeline run date; dataset snapshot per Hou et al. 2024)"],
            ["Fair-use note", "Academic coursework use per Amazon Reviews 2023 dataset terms and citation requirements"],
        ]
    )

    add_page_break(doc)

    # =========================================================================
    # Q2 — LLM ANALYSIS
    # =========================================================================
    heading(doc, "Q2 — LLM Analysis of Reviews", 1)

    heading(doc, "2.1  Overview and Goals", 2)
    body(doc,
        "The Q2 step transforms raw listing text and customer reviews into a structured, visually "
        "actionable analysis. The output is a typed ImagePromptBundle schema (src/models.py) that "
        "bridges text evidence to image-generation prompts for Q3. Concretely, the LLM must: "
        "(1) extract which visual cues reviews consistently describe; "
        "(2) distinguish consensus from outliers; "
        "(3) map which official listing bullets reviewers actually echo or dispute; and "
        "(4) produce 3–5 planned shots per product, each with a rationale_from_reviews and a full "
        "text-to-image prompt that respects diffusion model limitations."
    )

    heading(doc, "2.2  Prompting Strategy", 2)

    heading(doc, "Strategy 1 — Structured JSON output with chain-of-thought schema", 3)
    body(doc,
        "The LLM (gpt-5.4-mini, temperature 0.35) receives a system message positioning it as a "
        "creative director, enumerating critical constraints (no legible micro-text on seals, one "
        "scene per shot, brand inferences only from this product's text). The user message presents "
        "product metadata + RAG context followed by an explicit chain-of-thought instruction that "
        "orders the reasoning: (1) category norms, (2) consumer landscape, (3) listing↔review "
        "mapping, (4) cross-review benefits, (5) brand visual codes, (6) shot plan, (7) prompts. "
        "Output is forced to JSON via response_format: json_object and validated against a Pydantic "
        "schema. This ensures analytical traceability: every planned shot includes a "
        "rationale_from_reviews field explaining which review themes motivated it."
    )

    heading(doc, "Strategy 2 — Multi-query RAG augmentation", 3)
    body(doc,
        "Rather than stuffing the full review corpus into a single context window, we use Chroma "
        "(persistent vector store, data/chroma/) with text-embedding-3-small embeddings. Before "
        "calling the LLM, 12 targeted retrieval queries are issued in sequence; retrieved chunks "
        "are deduplicated by text key and concatenated as the RAG context block."
    )
    body(doc,
        "Why 12 queries instead of one? A single broad query surfaces the highest-rated or most "
        "embedding-similar chunks, which may all repeat the same dominant theme. A diverse query "
        "set — covering concerns, benefits, expectation gaps, visual/material properties, sensory "
        "outcomes, listing-review mapping, cross-review consensus, and brand/packaging mentions — "
        "forces the retriever to surface different facets of the review corpus."
    )

    add_table(doc,
        headers=["#", "RAG Retrieval Query"],
        rows=[
            ["Q1", "What concerns, complaints, risks, disappointments, defects, damage, or returns do customers mention?"],
            ["Q2", "What benefits, relief, positive outcomes, loyalty, or praise patterns appear in reviews?"],
            ["Q3", "What expectations versus reality show up: misleading photos, wrong size, wrong texture or type?"],
            ["Q4", "What do customers say about appearance, packaging, colour, size, shape, texture, sheen, matte or glossy finish?"],
            ["Q5", "What materials, fabric, foam, plastic, liquid, fill, weave, stitching, or construction details are described?"],
            ["Q6", "What defects, wear after washing, pilling, fading, or quality issues are mentioned?"],
            ["Q7", "What outcomes do reviews describe: cleaning results, efficacy, comfort, skin feel, sleep, fit?"],
            ["Q8", "What sensory or physical qualities appear: softness, firmness, weight, warmth, coolness, rigidity?"],
            ["Q9", "What mood, atmosphere, lifestyle, or setting do reviews imply?"],
            ["Q10", "Which features from the product listing do reviewers echo, praise, argue about, or say matter most?"],
            ["Q11", "Which benefits, outcomes, or feelings do multiple different reviewers repeat independently?"],
            ["Q12", "What do reviewers say about brand name, packaging colours, bottle or box design, logo area, or overall look?"],
        ]
    )

    heading(doc, "2.3  Chunking Strategy", 2)

    add_table(doc,
        headers=["Strategy", "Env var", "Behaviour"],
        rows=[
            ["Sliding window", "CHUNK_STRATEGY=sliding",
             "Splits each field into overlapping windows of 400 tokens with 60-token overlap between consecutive windows."],
            ["Paragraph batch (default)", "CHUNK_STRATEGY=paragraph_batch",
             "Splits on blank lines (\\n\\n); packs whole paragraphs up to 400 tokens; falls back to sliding only for oversized paragraphs."],
        ]
    )

    heading(doc, "Chunking A/B Results (Experiment 1)", 3)
    body(doc,
        "We ran both strategies on the identical three-product corpus (502 chunks indexed per "
        "product for both strategies), deleting data/chroma/ between runs. The metric is "
        "unique_chunks_used after deduplicating across all 12 RAG queries:"
    )

    add_table(doc,
        headers=["Strategy", "B0C4BFWLNC (pet bed)", "B0BG9Q18ZZ (CeraVe)", "B0BXSQN4HV (hat washer)", "Total indexed / product"],
        rows=[
            ["sliding", "46", "41", "43", "502"],
            ["paragraph_batch", "46", "42", "43", "502"],
        ]
    )

    body(doc,
        "Decision: We adopt paragraph_batch. Amazon reviews are predominantly short, self-contained "
        "paragraphs. Sliding windows over a short review create artificial overlapping chunks that "
        "repeat content with minor boundary offsets, adding retrieval noise without enriching context. "
        "The one-chunk difference on CeraVe (41 vs. 42) confirms the boundary effect is real but "
        "small. Archived experiment outputs are in outputs/chunking_ab/sliding/ and "
        "outputs/chunking_ab/paragraph_batch/."
    )

    heading(doc, "2.4  ImagePromptBundle Schema", 2)
    body(doc,
        "The ImagePromptBundle (src/models.py) is the structured bridge between Q2 and Q3. "
        "Key fields:"
    )

    add_table(doc,
        headers=["Field", "Purpose"],
        rows=[
            ["product_summary", "3–5 sentence blend of product facts and reviewer emotional tone"],
            ["category_imagery_norms", "Shot types that persuade shoppers in this category"],
            ["consumer_needs_and_concerns", "Recurring needs and complaints from reviews"],
            ["listing_features_elevated_by_reviews", "Official bullets that reviewers specifically echo or demand"],
            ["cross_review_benefits", "Benefits multiple reviewers independently repeat"],
            ["brand_visual_codes", "On-brand palette, finish, and mood cues inferred from title/description/reviews"],
            ["shot_plan_rationale", "Paragraph explaining the logic and ordering of planned shots"],
            ["pitfalls_to_avoid", "Customer-flagged issues + diffusion-safety constraints"],
            ["planned_shots", "3–5 PlannedImageShot objects (shot_index, role, rationale_from_reviews, prompt)"],
            ["pipeline_meta", "Injected by pipeline: retrieval queries, unique_chunks_used, model, chunking config"],
        ]
    )

    heading(doc, "Unique chunks used in final run (paragraph_batch, 12 queries)", 3)
    add_table(doc,
        headers=["Product", "ASIN", "unique_chunks_used"],
        rows=[
            ["Best Friends by Sheri donut bed", "B0C4BFWLNC", "67"],
            ["CeraVe Hydrating Facial Cleanser", "B0BG9Q18ZZ", "53"],
            ["YEENOR hat washer", "B0BXSQN4HV", "56"],
        ]
    )
    body(doc,
        "The pet bed draws the most unique chunks (67), consistent with its high review volume and "
        "the diversity of reviewer concerns (shape, texture, washability, calming behaviour, colour "
        "accuracy, size fit)."
    )

    add_page_break(doc)

    # =========================================================================
    # Q3 — IMAGE GENERATION
    # =========================================================================
    heading(doc, "Q3 — Image Generation with Diffusion Models", 1)

    heading(doc, "3.1  Models and Configuration", 2)
    add_table(doc,
        headers=["Backend", "Model", "Output directory", "Provider suffix"],
        rows=[
            ["OpenAI Images", "chatgpt-image-latest", "outputs/<ASIN>/openai/",
             '"Emphasize crisp commercial lighting and a catalog-ready composition."'],
            ["Google Gemini", "gemini-3.1-flash-image-preview", "outputs/<ASIN>/gemini/",
             '"Render as one coherent photographic scene; single clear focal subject."'],
        ]
    )
    body(doc,
        "Image size: 1024×1024 (OpenAI); 1K resolution (Gemini). Images generated: first 3 planned "
        "shots × 2 models × 3 products = 18 images total. Full configuration is recorded in "
        "outputs/run_log.json → models_used."
    )

    heading(doc, "3.2  Generated Images — Product 1: Best Friends by Sheri Donut Bed (B0C4BFWLNC)", 2)

    body(doc,
        "Shot 1 (hero_product): Reviewers care most about whether the bed looks truly plush and "
        "supportive, or flat and misleading. The hero foregrounds the raised rim, CoziLOFT loft, "
        "and Dark Chocolate faux-fur palette with the matching blanket visible."
    )
    side_by_side(doc,
        IMG["bed_hero_oai"], IMG["bed_hero_gem"],
        "OpenAI — chatgpt-image-latest", "Gemini — gemini-3.1-flash-image-preview",
        "Figure 1. Pet Bed — Shot 1 (hero_product). Left: OpenAI; Right: Gemini.",
    )

    body(doc,
        "Shot 2 (consumer_in_use): Reviewers repeatedly describe pets settling in and sleeping "
        "soundly — the key calming benefit. This shot demonstrates burrow-friendly comfort and the "
        "self-warming nest effect. Gemini's version adds warmer environmental context (living room "
        "setting) that enhances the lifestyle narrative."
    )
    side_by_side(doc,
        IMG["bed_use_oai"], IMG["bed_use_gem"],
        "OpenAI — chatgpt-image-latest", "Gemini — gemini-3.1-flash-image-preview",
        "Figure 2. Pet Bed — Shot 2 (consumer_in_use). Left: OpenAI; Right: Gemini.",
    )

    body(doc,
        "Shot 3 (wash_and_bundle_value): The blanket is a repeatedly praised separate utility item; "
        "this shot establishes bundle value and the travel/crate-cover use case."
    )
    side_by_side(doc,
        IMG["bed_bundle_oai"], IMG["bed_bundle_gem"],
        "OpenAI — chatgpt-image-latest", "Gemini — gemini-3.1-flash-image-preview",
        "Figure 3. Pet Bed — Shot 3 (wash_and_bundle_value). Left: OpenAI; Right: Gemini.",
    )

    if IMG["bed_detail_oai"].exists() and IMG["bed_detail_gem"].exists():
        body(doc,
            "Additional shot (construction_detail): Directly addresses the most repeated praise and "
            "criticism — loft, centre support, and whether the bed holds its shape after washing."
        )
        side_by_side(doc,
            IMG["bed_detail_oai"], IMG["bed_detail_gem"],
            "OpenAI — chatgpt-image-latest", "Gemini — gemini-3.1-flash-image-preview",
            "Figure 4. Pet Bed — Additional Shot (construction_detail). Left: OpenAI; Right: Gemini.",
        )

    heading(doc, "3.3  Generated Images — Product 2: CeraVe Hydrating Facial Cleanser (B0BG9Q18ZZ)", 2)

    body(doc,
        "Shot 1 (hero_product): This frame establishes the exact bottle people expect to receive "
        "and reassures shoppers worried about leaks, watered-down product, or pump issues. The "
        "clinical blue-and-white palette and satin-plastic finish reinforce the dermatologist-led "
        "brand_visual_codes."
    )
    side_by_side(doc,
        IMG["cer_hero_oai"], IMG["cer_hero_gem"],
        "OpenAI — chatgpt-image-latest", "Gemini — gemini-3.1-flash-image-preview",
        "Figure 5. CeraVe Cleanser — Shot 1 (hero_product). Left: OpenAI; Right: Gemini.",
    )

    body(doc,
        "Shot 2 (consumer_in_use): Reviewers repeatedly describe this cleanser as best understood "
        "on damp skin — a smooth lotion-like spread rather than foam. This shot highlights the "
        "cross-review benefit of skin feeling moisturised and not stripped, with calm bathroom "
        "staging. OpenAI's more controlled lighting better preserves the clinical, pharmacy-shelf "
        "mood that reviewers associate with the brand."
    )
    side_by_side(doc,
        IMG["cer_use_oai"], IMG["cer_use_gem"],
        "OpenAI — chatgpt-image-latest", "Gemini — gemini-3.1-flash-image-preview",
        "Figure 6. CeraVe Cleanser — Shot 2 (consumer_in_use). Left: OpenAI; Right: Gemini.",
    )

    body(doc,
        "Shot 3 (texture_formula_closeup): This frame directly answers the biggest expectation gap "
        "in reviews — the cleanser is not a bubbly foam but a lotion-like, smooth formula. It "
        "elevates the official hyaluronic acid, ceramides, and glycerin hydration story by showing "
        "a creamy texture cue without requiring legible ingredient text."
    )
    side_by_side(doc,
        IMG["cer_tex_oai"], IMG["cer_tex_gem"],
        "OpenAI — chatgpt-image-latest", "Gemini — gemini-3.1-flash-image-preview",
        "Figure 7. CeraVe Cleanser — Shot 3 (texture_formula_closeup). Left: OpenAI; Right: Gemini.",
    )

    heading(doc, "3.4  Generated Images — Product 3: YEENOR Hat Washer 2-Pack (B0BXSQN4HV)", 2)

    body(doc,
        "Shot 1 (hero_product): The hero makes the white 2-pack cage form instantly legible "
        "as a molded cap washer. Reviewers split most sharply on fit and sturdiness so shoppers "
        "need a truthful silhouette first. OpenAI's studio rendering more clearly shows the "
        "four-sided cage structure and clip geometry."
    )
    side_by_side(doc,
        IMG["hat_hero_oai"], IMG["hat_hero_gem"],
        "OpenAI — chatgpt-image-latest", "Gemini — gemini-3.1-flash-image-preview",
        "Figure 8. Hat Washer — Shot 1 (hero_product). Left: OpenAI; Right: Gemini.",
    )

    body(doc,
        "Shot 2 (consumer_in_use): Multiple reviewers specifically praise dishwasher use, easy "
        "operation, and hats coming out clean while retaining shape. Gemini's version adds more "
        "realistic appliance context that better supports this use case narrative."
    )
    side_by_side(doc,
        IMG["hat_use_oai"], IMG["hat_use_gem"],
        "OpenAI — chatgpt-image-latest", "Gemini — gemini-3.1-flash-image-preview",
        "Figure 9. Hat Washer — Shot 2 (consumer_in_use). Left: OpenAI; Right: Gemini.",
    )

    body(doc,
        "Shot 3 (detail_macro): The strongest praise centres on secure snaps and sturdier locking "
        "points while the main debate is fit and bill support. A macro detail isolates the 4-sided "
        "clip design, 10mm spacing, and rigid ribs that reviewers say keep hats secure."
    )
    side_by_side(doc,
        IMG["hat_mac_oai"], IMG["hat_mac_gem"],
        "OpenAI — chatgpt-image-latest", "Gemini — gemini-3.1-flash-image-preview",
        "Figure 10. Hat Washer — Shot 3 (detail_macro). Left: OpenAI; Right: Gemini.",
    )

    add_page_break(doc)

    # =========================================================================
    # MODEL COMPARISON (Experiment 2)
    # =========================================================================
    heading(doc, "3.5  Experiment 2 — Dual Image Model Comparison", 2)

    body(doc,
        "We route identical planned_shots[].prompt content through both chatgpt-image-latest and "
        "gemini-3.1-flash-image-preview, with only a short provider-specific suffix appended. "
        "This gives us a controlled comparison: prompt content is held constant, only the model changes."
    )

    add_table(doc,
        headers=["Dimension", "OpenAI (chatgpt-image-latest)", "Gemini (gemini-3.1-flash-image-preview)"],
        rows=[
            ["Composition", "Tighter studio/catalog frames; product centered and well-lit for e-commerce",
             "More likely to add lifestyle background (room context) even for product-focused prompts"],
            ["Colour accuracy", "Generally faithful to palette cues (dark chocolate fur, clinical white/blue)",
             "Sometimes richer and warmer in lifestyle mood; occasional saturation boost"],
            ["Text on product", "Micro-text (seals, barcodes) present but not actually readable at zoom",
             "Attempts more label detail; increases visible failure risk — text appears as plausible-but-wrong glyphs"],
            ["Physical plausibility", "Generally stable; rare floating or impossible staging",
             "Occasionally introduces extra objects or unusual framing not specified in the prompt"],
            ["Constraint adherence", "Follows 'no foam, no splash' instructions fairly reliably",
             "Sometimes generates text blocks or decorative elements explicitly excluded in the prompt"],
            ["Failure modes", "Seal silhouettes as soft generic shapes (as intended); occasional CGI over-polish",
             "Seal text attempts produce more visible hallucinated characters; lifestyle context can crowd the product"],
        ]
    )

    heading(doc, "Per-product model preferences", 3)
    add_table(doc,
        headers=["Product", "Preferred model", "Shot types", "Rationale"],
        rows=[
            ["Pet bed", "OpenAI", "Hero, construction detail",
             "Studio lighting captures dark-chocolate faux fur and raised rim more precisely"],
            ["Pet bed", "Gemini", "Consumer in-use (Shot 2)",
             "Added environmental warmth enhances the lifestyle/cozy narrative"],
            ["CeraVe cleanser", "OpenAI", "All three shots",
             "Clinical pharmacy-shelf aesthetic better served by catalog-frame tendency; Gemini softened trust-critical mood"],
            ["Hat washer", "OpenAI", "Hero, detail macro",
             "White plastic cage needs clear neutral studio light to make geometry legible"],
            ["Hat washer", "Gemini", "Consumer in-use (Shot 2)",
             "More realistic appliance context worked well for the in-dishwasher workflow"],
        ]
    )

    body(doc,
        "Overall insight: OpenAI chatgpt-image-latest is consistently safer for pack-shot clarity, "
        "controlled lighting, and constraint adherence — best for hero shots and detail macros. "
        "Gemini gemini-3.1-flash-image-preview has strengths in lifestyle mood and environmental "
        "richness when the scene requires consumer context. For a production listing, model "
        "selection per shot role is more effective than a single model for all frames."
    )

    add_page_break(doc)

    # =========================================================================
    # EXPERIMENT 3 — PROMPT ITERATION & REAL LISTING COMPARISON
    # =========================================================================
    heading(doc, "3.6  Experiment 3 — Real Listing Imagery Review and Prompt Iteration", 2)

    body(doc,
        "We loaded listing_image_urls from data/products.json (Amazon gallery order) for each "
        "ASIN. Real Amazon listings achieve high information density by spreading content across a "
        "carousel of specialised frames: a clean hero, a benefits/feature infographic, a size chart, "
        "a lifestyle image, and a detail/close-up shot. This observation directly shaped our "
        "multi-shot architecture: rather than generating one 'good AI image' per product, we plan "
        "3–5 shots that decompose the carousel story, each with one clear job."
    )

    heading(doc, "AI-generated vs. Real Listing Images — Comparison", 3)
    add_table(doc,
        headers=["Dimension", "Real Listing Photography", "AI Initial (pre-iteration)", "AI After Iteration"],
        rows=[
            ["SKU fidelity", "Exact product, exact colour, professionally photographed",
             "Generally faithful silhouette and palette",
             "Improved by anchoring prompts to specific listing features"],
            ["Label accuracy", "Pixel-perfect label text, certifications, regulatory print",
             "Seals and badges appear as plausible-but-wrong glyph clusters",
             "Substantially reduced by instructing models to omit/blur micro-text"],
            ["Information density", "Single frame carries one job well (hero, infographic, lifestyle)",
             "Initially overcrowded with multiple goals in one prompt",
             "Improved by strict one-scene-per-prompt discipline"],
            ["Physical plausibility", "Real-world staging, gravity-compliant",
             "Occasional floating, impossible balance, exaggerated proportions",
             "Improved by 'physically plausible staging' constraint"],
        ]
    )

    heading(doc, "Three rounds of prompt engineering changes (src/llm_analysis.py)", 3)
    body(doc,
        "Round 1 — RAG query expansion: Moved from narrow sensory queries to the 12-query hybrid "
        "set. Goal: surface a broader range of review evidence so the LLM has richer material.\n\n"
        "Round 2 — Structured fields for analytical traceability: Added "
        "listing_features_elevated_by_reviews, cross_review_benefits, brand_visual_codes, "
        "shot_plan_rationale, and pipeline_meta to the schema. Goal: make the reasoning chain "
        "explicit so each shot has a documented rationale.\n\n"
        "Round 3 — Diffusion-safe text and density constraints: After observing seal hallucinations "
        "and layout clutter, added explicit LLM instructions: (a) avoid legible micro-copy on seals "
        "and barcodes — use omit/blur/generic badge shape; (b) cap on-pack text to at most one or "
        "two short phrases from the Title/Official description; (c) one primary idea per prompt; "
        "(d) acknowledge no access to official brand DAM."
    )

    heading(doc, "What diffusion can and cannot do", 3)
    add_table(doc,
        headers=["Area", "AI Succeeds", "AI Struggles"],
        rows=[
            ["Silhouette & macro shape",
             "Donut bed round form, cleanser pump bottle, hat cage squared geometry — all recognisable",
             "Exact proportions and dimensions not guaranteed"],
            ["Colour palette & mood",
             "Dark chocolate fur, clinical blue-and-white, utilitarian white plastic — captured faithfully when prompted specifically",
             "Exact brand colour values cannot be specified numerically"],
            ["Lifestyle atmosphere",
             "Pet-in-bed comfort, bathroom sink cleansing, dishwasher staging — plausible and emotive",
             ""],
            ["Label text & certification seals",
             "",
             "Both models hallucinate seal text even with omit instructions; regulatory copy is beyond reliable t2i rendering"],
            ["Exact brand identity",
             "",
             "Without brand asset access, images infer a 'CeraVe-like' palette but cannot reproduce the exact logo/typeface"],
        ]
    )

    add_page_break(doc)

    # =========================================================================
    # Q4 — AI AGENTIC WORKFLOW
    # =========================================================================
    heading(doc, "Q4 — AI Agentic Workflow", 1)

    heading(doc, "4.1  Architecture Overview", 2)
    body(doc,
        "The pipeline is a genuine multi-agent system with four distinct agents coordinated by a "
        "Supervisor, typed JSON handoffs between agents, a structured harness with bounded retries "
        "and telemetry around each LLM call, and an explicit merge/verification step. It goes well "
        "beyond a single monolithic prompt.\n\n"
        "Key design principle: Q2 (Analyst) and Q3 (Creative Executive) are deliberately separated "
        "so the analyst role — grounded in reviews and data — never bleeds into the creative role "
        "— image-model-aware prompt writing. Each step has a single, auditable responsibility."
    )

    add_table(doc,
        headers=["Agent", "Role", "Primary artifact", "Implementation"],
        rows=[
            ["Q1 — Data/Catalog", "Offline data preparation: slice Amazon JSONL, consolidate products, apply rating filters",
             "data/products.json, data/products_manifest.json",
             "scripts/consolidate_products.py, scripts/filter_meta_by_rating.py"],
            ["Supervisor", "Orchestrates the full sequence; validates merge preconditions; logs every step",
             "outputs/run_log.json (AgentState)",
             "src/agent.py (run_pipeline())"],
            ["Q2 — Analyst", "Multi-query RAG over Chroma + structured LLM call; produces a review-grounded brief with shot roles and rationales — NO final image prompts",
             "ReviewImageryBrief JSON",
             "src/llm_analysis.py (Q2_SYSTEM, _build_q2_user())"],
            ["Q3 — Creative Executive", "Receives Q2 brief + listing metadata; produces final text-to-image prompts only — no re-analysis of reviews",
             "CreativePromptPack JSON",
             "src/llm_analysis.py (Q3_SYSTEM, _build_q3_user())"],
            ["Image Backends", "Optional rendering (not LLM agents): route each prompt through OpenAI Images and/or Gemini",
             "PNG files in outputs/<ASIN>/openai/ and .../gemini/",
             "src/image_gen.py"],
        ]
    )

    heading(doc, "4.2  End-to-End Pipeline Sequence", 2)
    add_table(doc,
        headers=["Stage", "Agent / Component", "What happens"],
        rows=[
            ["1. Load", "run_pipeline.py", "Load data/products.json → List[Product] (3 products)"],
            ["2. Chunk & Embed", "src/chunking.py → src/rag.py (ReviewRAGIndex)",
             "Product text → token-aware chunks (≤400 tokens, paragraph_batch default) → text-embedding-3-small → Chroma (data/chroma/); 502 chunks per product"],
            ["3. Retrieve", "RAG infrastructure",
             "12 queries issued sequentially → chunks deduplicated by text key → RAG context block (build_rag_context)"],
            ["4. Q2 Analyst", "StructuredLLMHarness('q2_analyst', OPENAI_TEXT_MODEL)",
             "Input: product metadata + RAG context. Output: ReviewImageryBrief — analysis fields + planned_shots with shot_index, role, rationale_from_reviews ONLY. Up to 2 attempts; temperature → 0.0 on retry. Telemetry: model, duration_ms, attempts, token counts."],
            ["5. Q3 Creative", "StructuredLLMHarness('q3_creative', OPENAI_Q3_TEXT_MODEL)",
             "Input: ReviewImageryBrief JSON + listing metadata (no raw reviews). Output: CreativePromptPack — {planned_shots: [{shot_index, prompt}]}. Must output exactly the same shot_index set as Q2. Up to 2 attempts; temperature → 0.0 on retry."],
            ["6. Supervisor merge", "merge_review_brief_with_creative() in src/models.py",
             "Asserts Q2 shot_index set == Q3 shot_index set (explicit failure on mismatch, no silent merge). Joins brief fields + creative prompts → ImagePromptBundle. Saves outputs/<ASIN>_image_prompt_bundle.json with full pipeline_meta including harness telemetry."],
            ["7. Image generation", "src/image_gen.py",
             "For each planned_shot (up to --images-per-model N): image_prompt_for_model() appends provider suffix → chatgpt-image-latest (OpenAI) + gemini-3.1-flash-image-preview (Gemini) → PNGs saved."],
            ["8. Log", "Supervisor (run_pipeline())",
             "outputs/run_log.json: all steps, models_used, analysis_meta with q2_harness and q3_harness telemetry dicts, image_paths."],
        ]
    )

    heading(doc, "4.3  StructuredLLMHarness — Reliability Infrastructure", 2)
    body(doc,
        "The StructuredLLMHarness (src/llm_harness.py) wraps every LLM call with the "
        "infrastructure that makes multi-agent systems reliable rather than brittle:"
    )
    add_table(doc,
        headers=["Practice", "Implementation"],
        rows=[
            ["Structured output", "response_format={\"type\": \"json_object\"} enforced at the API level"],
            ["Schema validation", "Pydantic model_validate() on every response — ReviewImageryBrief (Q2) and CreativePromptPack (Q3)"],
            ["Normalisation before validate", "_normalize_review_brief_json() and _normalize_creative_pack_json() coerce minor schema drift (e.g. string where list expected) before Pydantic sees the data"],
            ["Bounded retries", "Up to 2 attempts per call; temperature drops to 0.0 on retry for stability"],
            ["Loud failures", "RuntimeError raised (not silently swallowed) if all attempts fail"],
            ["Telemetry", "Each call returns {agent_id, model, attempts, duration_ms, prompt_tokens, completion_tokens, total_tokens} — merged into pipeline_meta.q2_harness / q3_harness in the run log"],
        ]
    )

    heading(doc, "4.4  Typed Handoffs Between Agents", 2)
    body(doc,
        "A key property of the agentic design is that each inter-agent handoff uses a typed JSON "
        "contract, not free text:"
    )
    add_table(doc,
        headers=["Handoff", "From → To", "Schema", "What is passed"],
        rows=[
            ["1", "Q1 → Supervisor", "Product (Pydantic)", "ASIN, title, description, reviews, listing_image_urls"],
            ["2", "RAG → Q2", "Numbered excerpt block", "Deduplicated review + description chunks"],
            ["3", "Q2 → Q3", "ReviewImageryBrief (JSON)", "All analysis fields + planned_shots with roles and rationales ONLY — no prompts"],
            ["4", "Q3 → Supervisor", "CreativePromptPack (JSON)", "{\"planned_shots\": [{shot_index, prompt}]} — the prompts only"],
            ["5", "Supervisor → disk / image step", "ImagePromptBundle (JSON)", "Merged output of Q2 + Q3; saved to outputs/<ASIN>_image_prompt_bundle.json"],
        ]
    )
    body(doc,
        "The merge step (merge_review_brief_with_creative()) explicitly asserts that the Q2 and "
        "Q3 shot_index sets are identical before combining them. Any mismatch raises an error — "
        "preventing silent data corruption if the creative agent hallucinated an extra shot or "
        "dropped one."
    )

    heading(doc, "4.5  Why Two LLM Calls Instead of One", 2)
    body(doc,
        "The separation of Q2 (analyst) and Q3 (creative) is a deliberate architectural choice:\n\n"
        "Q2 Analyst receives raw review text and listing descriptions. Its system prompt (Q2_SYSTEM) "
        "establishes a senior e-commerce analyst persona focused on evidence-grounding: what "
        "reviewers actually said, which features they elevate, what the consensus visual picture is. "
        "It explicitly does not write image-model prompts — its output contains only shot roles and "
        "rationales.\n\n"
        "Q3 Creative Executive never sees raw reviews. It receives only the structured Q2 brief "
        "and the listing title/description. Its system prompt (Q3_SYSTEM) establishes a creative "
        "director persona focused entirely on image-model constraints: avoiding legible micro-text, "
        "one scene per prompt, physically plausible staging, brand-inferred palette. By isolating "
        "this responsibility, we prevent analytical commentary from cluttering the creative prompt, "
        "and prevent the analyst from needing to know about image-model limitations.\n\n"
        "This separation also enables different models per stage: OPENAI_TEXT_MODEL for Q2 (factual, "
        "structured analysis) and OPENAI_Q3_TEXT_MODEL for Q3 (creative, longer-form prompt "
        "writing) — configurable independently in .env."
    )

    heading(doc, "4.6  Reproduce the Final Run", 2)
    body(doc,
        "cp .env.example .env   # fill in OPENAI_API_KEY, GEMINI_API_KEY\n"
        "cp data/products.example.json data/products.json\n"
        "python run_pipeline.py --images-per-model 3\n\n"
        "To reproduce the A/B chunking experiment:\n"
        "CHUNK_STRATEGY=sliding python run_pipeline.py --skip-images\n"
        "rm -rf data/chroma\n"
        "CHUNK_STRATEGY=paragraph_batch python run_pipeline.py --skip-images"
    )

    add_page_break(doc)

    # =========================================================================
    # REFLECTION
    # =========================================================================
    heading(doc, "Reflection — Challenges and Lessons Learned", 1)

    heading(doc, "Challenge 1: Direct Amazon scraping is not viable — pivot to an academic dataset", 2)
    body(doc,
        "Our initial plan was to scrape listing text and customer reviews directly from Amazon "
        "product pages using the target ASINs. Amazon's bot-detection systems blocked this "
        "consistently with 503 errors and CAPTCHA challenges, making live scraping an unreliable "
        "foundation for a reproducible academic pipeline. We pivoted to the Amazon Reviews 2023 "
        "dataset (Hou et al., 2024), which provides the same information — product metadata, listing "
        "descriptions, feature bullets, listing image URLs, and customer reviews — in pre-collected "
        "JSONL files per category.\n\n"
        "Lesson: For academic projects requiring reproducibility, pre-existing curated datasets are "
        "not a fallback — they are often the correct approach. Live scraping introduces fragility "
        "(rate limits, bot-detection, page-layout changes) that makes results hard to reproduce. "
        "Citing Hou et al. (2024) also gives the data collection step proper academic provenance.\n\n"
        "Limitation to acknowledge: The dataset is a snapshot (Amazon Reviews 2023) rather than a "
        "real-time pull. Review counts and listing content reflect the state of the catalog as of "
        "the dataset collection window. For a production system, a licensed data feed or the Amazon "
        "Product Advertising API would be the appropriate solution."
    )

    heading(doc, "Challenge 2: Diffusion text rendering is fundamentally unreliable for product labels", 2)
    body(doc,
        "The single biggest gap between AI-generated and real listing images is legible label text. "
        "Both models produce visually plausible label areas — the CeraVe bottle looks like it has a "
        "label — but zooming in reveals garbled, hallucinated glyphs. This is not a prompt engineering "
        "failure; it reflects a fundamental limitation of current diffusion models. Our mitigation "
        "(instruct the LLM to omit seals, blur badges, and limit on-pack text to short verified phrases) "
        "substantially reduced but did not eliminate the problem. For production-quality labeling, "
        "composite design or reference-conditioned generation (e.g., ControlNet with the actual label "
        "as reference) would be necessary.\n\n"
        "Lesson: Design the analysis step to explicitly categorise which visual claims can and cannot "
        "be reliably depicted by diffusion. The pitfalls_to_avoid field in ImagePromptBundle is the "
        "direct implementation of this lesson."
    )

    heading(doc, "Challenge 3: RAG retrieval quality depends on query diversity, not just volume", 2)
    body(doc,
        "Our initial single-query RAG approach retrieved chunks that were topically similar but not "
        "diverse — the top retrieved documents all discussed the same dominant theme (e.g., softness "
        "for the pet bed). Moving to a 12-query multi-query set dramatically improved coverage, "
        "surfacing 53–67 unique chunks per product (after deduplication) that cover visual, structural, "
        "emotional, and brand facets of the review corpus.\n\n"
        "Lesson: For short-text corpora (Amazon reviews are often ≤100 words), diverse retrieval "
        "queries outperform a single broad query."
    )

    heading(doc, "Challenge 4: Chunking strategy matters less than expected for this corpus", 2)
    body(doc,
        "We anticipated a meaningful difference between sliding and paragraph-batch chunking. In "
        "practice, for 500 reviews that are predominantly short (1–3 sentences), both strategies "
        "produce nearly identical retrieval breadth (max 1 unique-chunk difference). The real "
        "benefit of paragraph-batch is conceptual alignment: reviews are natural units of opinion, "
        "and packing whole reviews as chunks preserves that unit. Chunking strategy is more "
        "consequential when reviews are long narratives."
    )

    heading(doc, "Challenge 5: One prompt = one scene requires discipline against feature creep", 2)
    body(doc,
        "The natural instinct when writing image prompts is to include everything. A prompt packed "
        "with multiple primary subjects produces images where the AI hedges across all of them, "
        "delivering none well. Enforcing the one-scene, one-primary-idea discipline — formalised in "
        "the LLM's system instructions and in the schema's planned_shots decomposition — was the "
        "most impactful single intervention for image quality. The analogy to multi-slide listing "
        "carousels was useful: real listings don't cram a hero, infographic, and lifestyle shot into "
        "one frame."
    )

    heading(doc, "What we would do differently", 2)
    body(doc,
        "1. Use the Amazon Product Advertising API or a licensed feed for live, reproducible data "
        "pulls instead of relying on a dataset snapshot. This would allow real-time review collection "
        "tied to specific ASINs without scraping restrictions.\n\n"
        "2. Expand to 5 shots per product and both models for a richer Q3 comparison dataset "
        "(time constraint limited us to 3 shots per model per product).\n\n"
        "3. Add a baseline condition: run the LLM with no RAG (description only) and compare bundle "
        "quality to the RAG-augmented version — quantifying how much retrieval adds to visual specificity.\n\n"
        "4. Quantitative image evaluation: use CLIP similarity between generated images and real "
        "listing images, or structured human ratings on the four comparison dimensions, rather than "
        "purely qualitative commentary.\n\n"
        "5. Larger reviews corpus: for the pet bed (~68,000 ratings), 500 reviews may miss structured "
        "sub-populations by pet size, breed, or use case that a larger sample would surface."
    )

    add_page_break(doc)

    # =========================================================================
    # APPENDIX
    # =========================================================================
    heading(doc, "Appendix", 1)

    heading(doc, "A — Repository Layout", 2)
    body(doc,
        "Project/\n"
        "├── run_pipeline.py              # CLI entrypoint\n"
        "├── requirements.txt\n"
        "├── .env.example                 # Config template\n"
        "├── data/\n"
        "│   ├── products.json            # Three-product catalog\n"
        "│   └── chroma/                  # Vector index (gitignored)\n"
        "├── docs/\n"
        "│   ├── report_draft.md\n"
        "│   ├── report_final.md\n"
        "│   └── report_final.docx        # This document\n"
        "├── outputs/\n"
        "│   ├── run_log.json\n"
        "│   ├── B0C4BFWLNC_image_prompt_bundle.json\n"
        "│   ├── B0BG9Q18ZZ_image_prompt_bundle.json\n"
        "│   ├── B0BXSQN4HV_image_prompt_bundle.json\n"
        "│   ├── B0C4BFWLNC/openai/ & gemini/\n"
        "│   ├── B0BG9Q18ZZ/openai/ & gemini/\n"
        "│   ├── B0BXSQN4HV/openai/ & gemini/\n"
        "│   └── chunking_ab/sliding/ & paragraph_batch/\n"
        "└── src/\n"
        "    ├── agent.py          # Pipeline orchestration\n"
        "    ├── llm_analysis.py   # RAG queries + LLM prompts\n"
        "    ├── rag.py            # Chroma index + retrieval\n"
        "    ├── chunking.py       # sliding / paragraph_batch\n"
        "    ├── image_gen.py      # Image APIs + suffix routing\n"
        "    ├── models.py         # Pydantic schemas\n"
        "    └── config.py         # Env-loaded settings"
    )

    heading(doc, "B — Run Metadata (from outputs/run_log.json)", 2)
    add_table(doc,
        headers=["Key", "Value"],
        rows=[
            ["text model", "gpt-5.4-mini"],
            ["embedding model", "text-embedding-3-small"],
            ["openai_image_model", "chatgpt-image-latest"],
            ["gemini_image_model", "gemini-3.1-flash-image-preview"],
            ["image_size (OpenAI)", "1024×1024"],
            ["image_resolution (Gemini)", "1K"],
            ["openai_image_prompt_suffix", "Emphasize crisp commercial lighting and a catalog-ready composition."],
            ["gemini_image_prompt_suffix", "Render as one coherent photographic scene; single clear focal subject."],
            ["chunking strategy", "paragraph_batch"],
            ["CHUNK_MAX_TOKENS", "400"],
            ["CHUNK_OVERLAP_TOKENS", "60"],
        ]
    )

    heading(doc, "C — Deliverables Checklist", 2)
    for item in [
        "data/products.json — three-product catalog with listings, reviews, and image URLs",
        "src/ — full pipeline implementation (agent.py, llm_analysis.py, rag.py, chunking.py, image_gen.py, models.py, config.py)",
        "run_pipeline.py — CLI entrypoint with --skip-images and --images-per-model flags",
        "src/llm_analysis.py — system/user LLM prompt templates",
        "outputs/*_image_prompt_bundle.json — planned shot prompt values per ASIN",
        "outputs/<ASIN>/openai/*.png + outputs/<ASIN>/gemini/*.png — 18 generated images",
        "outputs/run_log.json — full run provenance",
        "outputs/chunking_ab/sliding/ and outputs/chunking_ab/paragraph_batch/ — A/B experiment archives",
        "docs/report_final.md + docs/report_final.docx — this report",
    ]:
        p = doc.add_paragraph(f"✓  {item}", style="List Bullet")
        p.runs[0].font.size = Pt(10)

    heading(doc, "D — References", 2)
    refs = [
        "Amazon Reviews 2023 dataset (McAuley Lab, UCSD). Pet Supplies, Beauty & Personal Care, and Home categories.",
        "OpenAI API — gpt-5.4-mini (text/JSON), text-embedding-3-small (embeddings), chatgpt-image-latest (images). platform.openai.com",
        "Google Gemini API — gemini-3.1-flash-image-preview (images). ai.google.dev",
        "Chroma — open-source embedding database. trychroma.com",
        "tiktoken — OpenAI tokenizer library (cl100k_base encoding).",
        "Pydantic v2 — data validation and schema definitions (src/models.py).",
        "94-844 Generative AI Lab Final Project specification (Spring 2026), Prof. Beibei Li, Heinz College, Carnegie Mellon University.",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph(f"[{i}] {r}")
        p.style = doc.styles["Normal"]
        p.runs[0].font.size = Pt(10)

    # =========================================================================
    # Save
    # =========================================================================
    out_path = DOCS / "report_final.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
